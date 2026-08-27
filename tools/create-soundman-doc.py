#!/usr/bin/env python3
"""Build a soundman-facing setlist run sheet from Jam song metadata."""

from __future__ import annotations

import json
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SONGS_DIR = ROOT / "songs"
OUT_DOCX = ROOT / "docs" / "soundman-setlist-show-2026-05-29.docx"
OUT_HTML = ROOT / "docs" / "soundman-setlist-show-2026-05-29.html"
OUT_STAGE_HTML = ROOT / "docs" / "stage-setlist-show-2026-05-29.html"
OUT_SONG_LIST_HTML = ROOT / "docs" / "song-list-show-2026-05-29.html"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "222222"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
NOTE_FILL = "FFF2CC"
BORDER = "B7C9DC"


MIX_GUIDANCE = {
    "tzar-li-charlie": (
        "Blues/funk opener. Keep drums, bass, and guitar tight at the start; "
        "bring trumpet forward for the early feature. Ronen has a solo; the "
        "guitar lift should cut without getting brittle."
    ),
    "ah-ah-ah": (
        "Clean/pop rhythm. Keep keys and vocal clear; note the quiet piano-only "
        "moments and clapping section before the repeated final choruses."
    ),
    "ahava-hadasha": (
        "Bass and guitar riff start. Guitar target is clean/chorus-friendly; "
        "watch for excessive brightness or hard upper mids."
    ),
    "vampire": (
        "Piano-led ballad with edge-of-breakup guitar. Preserve vocal intimacy; "
        "let bigger sections grow without washing out the piano."
    ),
    "these-boots": (
        "Twangy clean guitar with bass-walk feel. Ronen has a guitar solo with "
        "a boost; check that it lifts without harsh treble."
    ),
    "parperei-titua": (
        "Ronen mainly enters for chorus/elevation, then ending solo. Watch for "
        "phaser on the ending solo and keep it present but not smeared."
    ),
    "shir-hamakolet": (
        "Clean/crunch rhythm. Compressor is intended to add body; keep vocals "
        "and groove intelligible rather than pushing guitar gain."
    ),
    "valerie": (
        "Compressed funk/soul pocket. Keep bass and drums tight, guitar clean "
        "and percussive, and leave room for the lead vocal."
    ),
    "mishehu": (
        "Expressive build. The second part needs a bigger guitar boost and more "
        "drama; prefer width/mids over extra harsh gain."
    ),
    "blues-cnaani": (
        "Blues-rock target with solos. Use one guitar sound for rhythm edge and "
        "one for lift; avoid over-compressing the groove."
    ),
    "yehudim-kah-oti": (
        "Dirty rhythm with reverb/slight delay. End section has solo 1 bigger, "
        "then solo 2 with wah for the screaming lead; watch treble peaks."
    ),
    "all-the-small-things": (
        "Tight pop-punk. Marketplace Blink 182 tone with added room/reverb; "
        "keep it less dry but do not soften palm-muted rhythm."
    ),
    "kerach-9": (
        "Fast rhythm crunch. Compressor adds body; if the guitar feels too "
        "polite, add presence in the mix before asking for more gain."
    ),
    "nitzotzot": (
        "Moderate drive with a solo. Two-drive setup should give edge plus lift; "
        "make the solo about sustain and mids, not just saturation."
    ),
    "haperach-begani": (
        "Wah-friendly funky rhythm. Keep bass tight and treble controlled; too "
        "much pre-wah gain will get sharp."
    ),
}


PERSONNEL = {
    "tzar-li-charlie": {
        "singers": "Moria, Ronen",
        "solos": "Trumpet, Keyboard, Guitar, Kazu",
    },
    "ah-ah-ah": {"singers": "Moria"},
    "ahava-hadasha": {"singers": "Ron", "solos": "Trumpet"},
    "vampire": {"singers": "Netta"},
    "these-boots": {"singers": "Netta", "solos": "Trumpet, Keyboard, Guitar"},
    "parperei-titua": {"singers": "Ronen", "solos": "Guitar"},
    "shir-hamakolet": {"singers": "Ronen, Ron, Ortal", "solos": "Keyboard"},
    "valerie": {"singers": "Moria"},
    "mishehu": {"singers": "Ortal"},
    "blues-cnaani": {
        "singers": "Moria, Ortal",
        "note": "Nadav acoustic gtr; Ron drums; Vardit on accordion; Ronen on tambourine",
    },
    "yehudim-kah-oti": {"singers": "Moria", "solos": "Guitar"},
    "all-the-small-things": {"singers": "Netta, Ron, Vardit"},
    "kerach-9": {"singers": "Moria"},
    "nitzotzot": {"singers": "Moria, Ronen", "solos": "Guitar, Trumpet"},
    "haperach-begani": {"singers": "Moria", "solos": "Trumpet", "note": "Encore"},
}


SOUNDMAN_BRIEF = {
    "tzar-li-charlie": "Blues/funk opener; trumpet forward early, then keyboard/guitar/Kazu features.",
    "ah-ah-ah": "Clean pop rhythm; protect vocal and quiet piano/clap moments.",
    "ahava-hadasha": "Bass/guitar riff start; trumpet solo feature.",
    "vampire": "Piano-led ballad; keep vocal intimate, let choruses grow.",
    "these-boots": "Twangy clean feel; trumpet, keyboard, and guitar solo lifts.",
    "parperei-titua": "Ronen enters for chorus/elevation; guitar solo with phaser color.",
    "shir-hamakolet": "Three singers; keyboard solo, keep groove/vocals clear.",
    "valerie": "Compressed funk/soul pocket; leave room for Moria lead.",
    "mishehu": "Expressive build for Ortal; second part needs bigger guitar lift.",
    "blues-cnaani": "Moria/Ortal vocals; Nadav acoustic guitar, Ron drums, Vardit on accordion, Ronen on tambourine.",
    "yehudim-kah-oti": "Dirty rhythm; guitar solo/wah end needs treble control.",
    "all-the-small-things": "Tight pop-punk; three singers, keep rhythm punchy.",
    "kerach-9": "Fast rhythm crunch; keep Moria vocal on top.",
    "nitzotzot": "Drive song with guitar and trumpet solos; solo boosts need mids.",
    "haperach-begani": "Encore. Funky wah rhythm; trumpet solo, controlled treble.",
}


OPENING_OVERRIDES = {
    "all-the-small-things": (
        "Band count-in to chart intro: C - F - G - G(F). Source opening cue is "
        "marked for verification."
    ),
}

ENDING_OVERRIDES = {
    "tzar-li-charlie": "End on Dm after the final chorus repeat.",
    "ah-ah-ah": "Final chorus repeats four times; end on piano-only C.",
    "ahava-hadasha": "Ending pattern: Em - G - F# - Em x4.",
    "these-boots": "Outro is bass-walk fade.",
    "parperei-titua": "Ending progression: C - Cmaj7 - F - A x2.",
    "valerie": "Outro fades on Db.",
    "mishehu": "Return to Bb; cadence on Bbmaj7.",
    "all-the-small-things": "Ending hits F - C.",
    "kerach-9": "Last chorus after the A/E modulation outro.",
}


SOURCE_CAVEATS = (
    "Several songs have blank BPM, key, or ending fields in the chart source. "
    "Use the band count-in where tempo is blank. The All The Small Things "
    "opening cue in source metadata appears stale, so this sheet uses the chart "
    "intro instead. Fender GTX presets are expected to follow show order; "
    "confirm device slots before the downbeat if older amp notes are in use."
)


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def load_show_songs() -> tuple[str, list[dict]]:
    index = json.loads((SONGS_DIR / "index.json").read_text(encoding="utf-8"))
    songs = []
    for entry in index["songs"]:
        if not entry.get("show"):
            continue
        data = json.loads((SONGS_DIR / f"{entry['slug']}.json").read_text(encoding="utf-8"))
        data["slug"] = entry["slug"]
        data["show"] = entry["show"]
        songs.append(data)
    songs.sort(key=lambda item: item["show"])
    return index.get("setlist", "Show"), songs


def set_style_font(style, name: str, size_pt: float | None = None, color: str | None = None):
    font = style.font
    font.name = name
    if size_pt:
        font.size = Pt(size_pt)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    set_style_font(styles["Normal"], "Arial", 9.5, INK)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    set_style_font(styles["Title"], "Arial", 22, INK)
    styles["Title"].font.bold = True
    styles["Title"].paragraph_format.space_after = Pt(2)

    set_style_font(styles["Subtitle"], "Arial", 10.5, MUTED)
    styles["Subtitle"].paragraph_format.space_after = Pt(10)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 11, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 3),
    ):
        set_style_font(styles[name], "Arial", size, color)
        styles[name].font.bold = True
        styles[name].paragraph_format.space_before = Pt(before)
        styles[name].paragraph_format.space_after = Pt(after)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width))


def set_table_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tc_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_in: list[float], indent: int = 120) -> None:
    widths = [dxa(value) for value in widths_in]
    total = sum(widths)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_cell_margins(table)
    set_table_borders(table)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.clear()


def write_cell(cell, text: str, *, size=8.5, bold=False, color=INK, fill=None, align=None) -> None:
    if fill:
        set_cell_shading(cell, fill)
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Arial")


def add_label_table(doc: Document, rows: list[tuple[str, str]], fill: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1.55, 4.95])
    repeat_header(table.rows[0])
    write_cell(
        table.rows[0].cells[0],
        "Item",
        size=8.2,
        bold=True,
        color=INK,
        fill=LIGHT_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    write_cell(
        table.rows[0].cells[1],
        "Detail",
        size=8.2,
        bold=True,
        color=INK,
        fill=LIGHT_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for label, detail in rows:
        cells = table.add_row().cells
        write_cell(cells[0], label, size=8.5, bold=True, color=DARK_BLUE, fill=fill or LIGHT_GRAY)
        write_cell(cells[1], detail, size=8.7, fill=fill)
    doc.add_paragraph()


def facts(song: dict) -> str:
    parts = []
    if song.get("key"):
        parts.append(str(song["key"]))
    if song.get("bpm"):
        parts.append(f"{song['bpm']} BPM")
    if song.get("timeSig"):
        parts.append(str(song["timeSig"]))
    return " / ".join(parts) if parts else "Confirm live"


def first_section_progression(song: dict) -> str:
    sections = song.get("sections") or []
    if not sections:
        return ""
    lines = sections[0].get("lines") or []
    if not lines:
        return ""
    first = lines[0]
    if isinstance(first, list):
        return " - ".join(str(item) for item in first)
    return ""


def opening(song: dict) -> str:
    if song["slug"] in OPENING_OVERRIDES:
        return OPENING_OVERRIDES[song["slug"]]
    if song.get("opening"):
        return str(song["opening"])
    if song.get("formSteps"):
        return f"Starts with: {song['formSteps'][0]}"
    progression = first_section_progression(song)
    if progression:
        return f"Starts with chart intro: {progression}"
    return "Band count-in / confirm live."


def ending(song: dict) -> str:
    if song["slug"] in ENDING_OVERRIDES:
        return ENDING_OVERRIDES[song["slug"]]
    if song.get("ending"):
        return str(song["ending"])
    if song.get("formSteps"):
        return f"Last landmark: {song['formSteps'][-1]}"
    return "Confirm ending live."


def form_brief(song: dict) -> str:
    steps = song.get("formSteps") or []
    if not steps:
        return ""
    if len(steps) <= 6:
        return " > ".join(steps)
    return " > ".join(steps[:3]) + " > ... > " + " > ".join(steps[-2:])


def song_label(song: dict) -> str:
    return f"{song['title']}\n{song.get('artist') or ''}".strip()


def personnel(song: dict, field: str) -> str:
    return PERSONNEL.get(song["slug"], {}).get(field, "")


def personnel_summary(song: dict) -> str:
    info = PERSONNEL.get(song["slug"], {})
    parts = []
    if info.get("singers"):
        parts.append(f"Singers: {info['singers']}")
    if info.get("solos"):
        parts.append(f"Solos: {info['solos']}")
    if info.get("note"):
        parts.append(f"Note: {info['note']}")
    return "; ".join(parts)


def cue_summary(song: dict) -> str:
    return f"Start: {opening(song)}\nEnd: {ending(song)}"


def add_run_sheet(doc: Document, songs: list[dict]) -> None:
    doc.add_heading("Setlist Run Sheet", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [0.38, 1.48, 0.88, 1.7, 2.06]
    set_table_geometry(table, widths)
    headers = ["#", "Song", "Key / Tempo", "Start + Form Anchor", "Sound Guidance / Ending Watch"]
    repeat_header(table.rows[0])
    for idx, text in enumerate(headers):
        write_cell(
            table.rows[0].cells[idx],
            text,
            size=8.2,
            bold=True,
            color=INK,
            fill=LIGHT_BLUE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for song in songs:
        cells = table.add_row().cells
        write_cell(cells[0], str(song["show"]), size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(cells[1], song_label(song), size=8.0, bold=True)
        write_cell(cells[2], facts(song), size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER)

        cue = opening(song)
        form = form_brief(song)
        if form:
            cue = f"{cue}\nForm: {form}"
        write_cell(cells[3], cue, size=7.6)

        guide = MIX_GUIDANCE.get(song["slug"], "Confirm mix notes live.")
        watch = ending(song)
        people = personnel_summary(song)
        prefix = f"{people}\n" if people else ""
        write_cell(cells[4], f"{prefix}{guide}\nEnding/watch: {watch}", size=7.4)


def add_quick_reference(doc: Document, songs: list[dict], setlist: str) -> None:
    title = doc.add_paragraph(style="Title")
    title.add_run(f"Soundman Run Sheet - Show {setlist}")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "15-song setlist summary for live mix, guitar/FX watch points, and arrangement cues."
    )

    add_label_table(
        doc,
        [
            (
                "Source",
                "Generated from songs/index.json, per-song JSON chart metadata, and amp-presets-show-2026-05-29.md.",
            ),
            (
                "Guitar rig",
                "Fender Mustang GTX 100, Fender American Professional II, GTX global EQ noted as Bright Cut.",
            ),
            (
                "Preset flow",
                "Use the show order as the intended GTX preset order: slots 1-15 map to songs 1-15.",
            ),
            (
                "Mix priority",
                "Keep vocals readable, level-match guitar presets, and treat lead boosts as mids/body/sustain rather than only extra volume.",
            ),
        ],
    )

    doc.add_heading("Fast Checks Before Downbeat", level=1)
    add_label_table(
        doc,
        [
            ("Source caveats", SOURCE_CAVEATS),
            (
                "Moments to mark",
                "Charlie, Boots, Mishehu, Kach Oti, Nitzotzot, and HaPerach have the most explicit guitar/FX watch points.",
            ),
            (
                "Tempo blanks",
                "Only fixed BPM values in source: Vampire 69, Boots 92, Valerie 102, All The Small Things 148, Kerach 9 160.",
            ),
        ],
        fill=NOTE_FILL,
    )

    doc.add_heading("GTX Preset Map", level=1)
    rows = []
    for song in songs:
        rows.append((f"{song['show']:02d}", f"{song['title']} - {MIX_GUIDANCE.get(song['slug'], '')}"))
    add_label_table(doc, rows)


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Jam - Show 5/29/26 soundman run sheet")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def html_text(value: str) -> str:
    return escape(value or "").replace("\n", "<br>")


def clean_html(value: str) -> str:
    """Remove template indentation from otherwise blank line endings."""
    return "\n".join(line.rstrip() for line in value.splitlines()) + "\n"


def build_html(setlist: str, songs: list[dict]) -> str:
    rows = []
    for song in songs:
        rows.append(
            f"""
            <tr>
              <td class="num">{song['show']}</td>
              <td class="song"><strong dir="auto">{html_text(song['title'])}</strong></td>
              <td>{html_text(personnel(song, 'singers') or 'Confirm')}</td>
              <td>{html_text(personnel(song, 'solos') or '-')}{('<br><strong>Note:</strong> ' + html_text(personnel(song, 'note'))) if personnel(song, 'note') else ''}</td>
              <td class="facts">{html_text(facts(song))}</td>
              <td>{html_text(cue_summary(song))}</td>
              <td>{html_text(SOUNDMAN_BRIEF.get(song['slug'], MIX_GUIDANCE.get(song['slug'], 'Confirm mix notes live.')))}</td>
            </tr>
            """
        )

    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Soundman Run Sheet - Show {html_text(setlist)}</title>
<style>
@page {{
  size: Letter landscape;
  margin: 0.28in 0.32in 0.28in 0.32in;
  @bottom-center {{
    content: "Jam - Show 5/29/26 soundman sheet - " counter(page) " / " counter(pages);
    color: #666;
    font-size: 7pt;
  }}
}}
* {{ box-sizing: border-box; }}
body {{
  margin: 0;
  color: #{INK};
  font-family: Arial, Helvetica, sans-serif;
  font-size: 7.1pt;
  line-height: 1.13;
}}
h1 {{
  margin: 0 0 0.035in 0;
  font-size: 18pt;
  line-height: 1.05;
  color: #{INK};
}}
h2 {{
  margin: 0.07in 0 0.035in 0;
  font-size: 10pt;
  color: #{BLUE};
}}
p {{ margin: 0 0 0.035in 0; }}
.subtitle {{
  color: #{MUTED};
  font-size: 7.8pt;
  margin-bottom: 0.05in;
}}
.meta-grid {{
  display: grid;
  grid-template-columns: 1fr 1fr 1fr;
  gap: 0.05in;
  margin: 0.05in 0 0.05in 0;
}}
.box {{
  border: 1px solid #{BORDER};
  background: #F8FAFC;
  padding: 0.045in 0.055in;
  min-height: 0.34in;
}}
.box strong {{
  color: #{DARK_BLUE};
  display: block;
  margin-bottom: 0.01in;
}}
.warn {{
  border-left: 5px solid #B7791F;
  background: #{NOTE_FILL};
  padding: 0.045in 0.06in;
  margin: 0.04in 0 0.055in 0;
}}
table {{
  border-collapse: collapse;
  width: 100%;
  table-layout: fixed;
  page-break-inside: auto;
}}
th, td {{
  border: 1px solid #{BORDER};
  padding: 0.025in 0.035in;
  vertical-align: top;
  overflow-wrap: anywhere;
}}
th {{
  background: #{LIGHT_BLUE};
  color: #{INK};
  font-weight: 700;
  text-align: center;
}}
tr {{ break-inside: avoid; page-break-inside: avoid; }}
.num {{ width: 0.32in; text-align: center; font-weight: 700; }}
.song {{ width: 1.18in; font-weight: 700; }}
.facts {{ width: 0.72in; text-align: center; }}
.singers {{ width: 0.9in; }}
.features {{ width: 1.3in; }}
.cue {{ width: 2.25in; }}
.guide {{ width: 3.45in; }}
strong {{ font-weight: 700; }}
</style>
</head>
<body>
  <h1>Soundman Run Sheet - Show {html_text(setlist)}</h1>
  <p class="subtitle">One-page mix run sheet: vocals, solos/features, start/end cues, and the main sound watch for each song.</p>
  <div class="meta-grid">
    <div class="box"><strong>Rig</strong>Fender Mustang GTX 100. Presets should follow slots 1-15 in show order.</div>
    <div class="box"><strong>Priority</strong>Lead vocals first, then featured solos. Guitar boosts should add mids/body/sustain, not just volume.</div>
    <div class="box"><strong>Tempo blanks</strong>Use band count-in where BPM is blank. Fixed: Vampire 69, Boots 92, Valerie 102, Blink 148, Kerach 160.</div>
  </div>
  <div class="warn"><strong>Before downbeat:</strong> Confirm GTX slot order, lead vocal mics, featured solo channels, and the Blues Cnaani special setup: Nadav acoustic guitar, Ron drums, Vardit on accordion, Ronen on tambourine.</div>

  <h2>Setlist</h2>
  <table>
    <colgroup>
      <col style="width:0.32in">
      <col style="width:1.18in">
      <col style="width:0.9in">
      <col style="width:1.3in">
      <col style="width:0.72in">
      <col style="width:2.25in">
      <col style="width:3.45in">
    </colgroup>
    <thead>
      <tr>
        <th>#</th>
        <th>Song</th>
        <th>Singers</th>
        <th>Solos / Notes</th>
        <th>Key / Tempo</th>
        <th>Start / End</th>
        <th>Sound Watch</th>
      </tr>
    </thead>
    <tbody>
      {''.join(rows)}
    </tbody>
  </table>
</body>
</html>
"""


def build_stage_html(setlist: str, songs: list[dict]) -> str:
    rows = []
    for song in songs:
        note_parts = []
        if personnel(song, "solos"):
            note_parts.append(f"Solos: {personnel(song, 'solos')}")
        if personnel(song, "note"):
            note_parts.append(personnel(song, "note"))
        rows.append(
            f"""
            <tr>
              <td class="num">{song['show']}</td>
              <td class="song" dir="auto">{html_text(song['title'])}</td>
              <td class="singers">{html_text(personnel(song, 'singers') or 'Confirm')}</td>
              <td class="notes">{html_text(' | '.join(note_parts) or '-')}</td>
            </tr>
            """
        )

    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Stage Setlist - Show {html_text(setlist)}</title>
<style>
@page {{
  size: Letter landscape;
  margin: 0.16in 0.2in;
}}
* {{ box-sizing: border-box; }}
html, body {{
  height: 8.18in;
}}
body {{
  margin: 0;
  color: #{INK};
  font-family: Arial, Helvetica, sans-serif;
}}
table {{
  width: 100%;
  height: 8.18in;
  border-collapse: collapse;
  table-layout: fixed;
}}
th, td {{
  border: 1.7pt solid #6F849C;
  padding: 0.055in 0.08in;
  vertical-align: middle;
  overflow-wrap: anywhere;
}}
th {{
  background: #{LIGHT_BLUE};
  font-size: 16pt;
  text-transform: uppercase;
  letter-spacing: 0;
}}
td {{
  font-size: 17.2pt;
  line-height: 1.08;
}}
tbody tr {{ height: 0.488in; }}
tr:nth-child(even) td {{ background: #F6F8FB; }}
.num {{ width: 0.56in; text-align: center; font-weight: 800; font-size: 18pt; }}
.song {{ width: 3.0in; font-weight: 800; }}
.singers {{ width: 2.18in; font-weight: 700; }}
.notes {{ width: 4.86in; font-size: 16.2pt; }}
</style>
</head>
<body>
  <table>
    <colgroup>
      <col style="width:0.56in">
      <col style="width:3.0in">
      <col style="width:2.18in">
      <col style="width:4.86in">
    </colgroup>
    <thead>
      <tr><th>#</th><th>Song</th><th>Singers</th><th>Solos / Stage Notes</th></tr>
    </thead>
    <tbody>
      {''.join(rows)}
    </tbody>
  </table>
</body>
</html>
"""


def build_song_list_html(setlist: str, songs: list[dict]) -> str:
    rows = []
    for song in songs:
        rows.append(
            f"""
            <tr>
              <td class="num">{song['show']}</td>
              <td class="song" dir="auto">{html_text(song['title'])}</td>
            </tr>
            """
        )

    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Song List - Show {html_text(setlist)}</title>
<style>
@page {{
  size: Letter landscape;
  margin: 0.16in 0.2in;
}}
* {{ box-sizing: border-box; }}
html, body {{
  height: 8.18in;
}}
body {{
  margin: 0;
  color: #{INK};
  font-family: Arial, Helvetica, sans-serif;
}}
table {{
  width: 100%;
  height: 8.18in;
  border-collapse: collapse;
  table-layout: fixed;
}}
td {{
  border: 1.9pt solid #6F849C;
  padding: 0.035in 0.085in;
  vertical-align: middle;
  overflow-wrap: anywhere;
}}
td {{
  font-size: 25.5pt;
  line-height: 1.0;
}}
tbody tr {{ height: 0.535in; }}
tr:nth-child(even) td {{ background: #F6F8FB; }}
.num {{ width: 0.78in; text-align: center; font-weight: 800; font-size: 23.5pt; }}
.song {{ width: 9.82in; font-weight: 800; }}
</style>
</head>
<body>
  <table>
    <colgroup>
      <col style="width:0.78in">
      <col style="width:9.82in">
    </colgroup>
    <tbody>
      {''.join(rows)}
    </tbody>
  </table>
</body>
</html>
"""


def build() -> Path:
    setlist, songs = load_show_songs()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    configure_styles(doc)
    add_quick_reference(doc, songs, setlist)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_footer(doc.sections[-1])
    add_run_sheet(doc, songs)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    OUT_HTML.write_text(clean_html(build_html(setlist, songs)), encoding="utf-8")
    OUT_STAGE_HTML.write_text(clean_html(build_stage_html(setlist, songs)), encoding="utf-8")
    OUT_SONG_LIST_HTML.write_text(
        clean_html(build_song_list_html(setlist, songs)), encoding="utf-8"
    )
    return OUT_DOCX


if __name__ == "__main__":
    print(build())
